import os
import openai
from app.config import settings
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
from pathlib import Path
from datetime import datetime
import mimetypes

# File size limit in bytes (50MB)
MAX_FILE_SIZE = 50 * 1024 * 1024

# Supported file types
SUPPORTED_FILE_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx"
}

# MIME type to extension mapping
MIME_TO_EXTENSION = {
    mime: ext for ext, mime in mimetypes.types_map.items()
}

# Prefer environment variable, fallback to settings from .env
openai.api_key = os.getenv("OPENAI_API_KEY") or settings.OPENAI_API_KEY

class DocumentProcessor:
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """
        Validate file size and type before processing.
        Returns a dictionary with validation results.
        """
        try:
            file_size = os.path.getsize(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            # Check file size
            if file_size > MAX_FILE_SIZE:
                return {
                    "valid": False,
                    "error": f"File size exceeds maximum limit of {MAX_FILE_SIZE / (1024*1024)} MB"
                }

            # Get MIME type using built-in mimetypes
            mime_type, _ = mimetypes.guess_type(file_path)
            
            if not mime_type:
                return {
                    "valid": False,
                    "error": f"Could not determine MIME type for file: {file_path}"
                }

            # Check if MIME type is supported
            if mime_type not in SUPPORTED_FILE_TYPES:
                return {
                    "valid": False,
                    "error": f"Unsupported file type: {mime_type}"
                }

            # Additional validation for specific file types
            if mime_type == "application/pdf":
                try:
                    PdfReader(file_path)
                except Exception as e:
                    return {
                        "valid": False,
                        "error": f"Invalid PDF file: {str(e)}"
                    }
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                try:
                    Document(file_path)
                except Exception as e:
                    return {
                        "valid": False,
                        "error": f"Invalid DOCX file: {str(e)}"
                    }
            elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                try:
                    load_workbook(file_path)
                except Exception as e:
                    return {
                        "valid": False,
                        "error": f"Invalid XLSX file: {str(e)}"
                    }

            return {
                "valid": True,
                "mime_type": mime_type,
                "size": file_size,
                "extension": file_ext,
                "file_path": file_path
            }

        except Exception as e:
            return {
                "valid": False,
                "error": f"Error validating file: {str(e)}"
            }

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from document after validation.
        Returns a dictionary containing extracted text and validation info.
        """
        validation = self.validate_file(file_path)
        if not validation["valid"]:
            return validation

        try:
            if validation["mime_type"] == "application/pdf":
                reader = PdfReader(file_path)
                text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif validation["mime_type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif validation["mime_type"] == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
                wb = load_workbook(file_path)
                text = ""
                for sheet in wb:
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " ".join([str(cell) for cell in row if cell])
                        text += row_text + "\n"
                text = text.strip()
            else:
                raise ValueError(f"Unexpected MIME type: {validation['mime_type']}")

            return {
                "valid": True,
                "text": text,
                "mime_type": validation["mime_type"],
                "size": validation["size"],
                "processed_at": datetime.utcnow().isoformat()
            }

        except Exception as e:
            return {
                "valid": False,
                "error": f"Error processing file: {str(e)}"
            }

    def analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """
        Analyze text sentiment using OpenAI API.
        Returns a structured sentiment analysis with score and categories.
        """
        if not text.strip():
            return {
                "valid": False,
                "error": "No text to analyze"
            }

        try:
            prompt = f"""
            Analyze the following text and provide:
            1. Overall sentiment (Positive, Neutral, Negative)
            2. Sentiment score between -1 (very negative) and 1 (very positive)
            3. Key topics mentioned
            4. Risk factors identified
            5. Business impact assessment

            Text:
            {text}
            """

            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )

            # Parse the response into a structured format
            result = response.choices[0].message.content.strip()
            
            # The model will return JSON-like text, parse it
            try:
                import json
                data = json.loads(result)
                return {
                    "valid": True,
                    "sentiment": data.get("sentiment", "neutral"),
                    "score": data.get("score", 0.0),
                    "key_topics": data.get("key_topics", []),
                    "risk_factors": data.get("risk_factors", []),
                    "business_impact": data.get("business_impact", ""),
                    "processed_at": datetime.utcnow().isoformat()
                }
            except json.JSONDecodeError:
                return {
                    "valid": False,
                    "error": "Failed to parse AI response"
                }

        except Exception as e:
            return {
                "valid": False,
                "error": f"Error analyzing sentiment: {str(e)}"
            }

    def extract_action_items(self, text: str) -> Dict[str, Any]:
        """
        Extract actionable items from text using OpenAI API.
        Returns a structured list of action items with metadata.
        """
        if not text.strip():
            return {
                "valid": False,
                "error": "No text to analyze"
            }

        try:
            prompt = f"""
            Extract all actionable items from this text. For each item, provide:
            - description: A clear description of the action
            - priority: high, medium, or low
            - responsible_party: Who should do this (if mentioned)
            - due_date: When it should be completed (if mentioned)
            - status: pending, in_progress, completed
            - dependencies: Other tasks this depends on (if any)
            - risk_level: high, medium, low

            Format the response as JSON array of objects.

            Text:
            {text}
            """

            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )

            # Parse the response into a structured format
            result = response.choices[0].message.content.strip()
            
            try:
                import json
                data = json.loads(result)
                return {
                    "valid": True,
                    "action_items": data,
                    "total_items": len(data),
                    "processed_at": datetime.utcnow().isoformat()
                }
            except json.JSONDecodeError:
                return {
                    "valid": False,
                    "error": "Failed to parse AI response"
                }

        except Exception as e:
            return {
                "valid": False,
                "error": f"Error extracting action items: {str(e)}"
            }
